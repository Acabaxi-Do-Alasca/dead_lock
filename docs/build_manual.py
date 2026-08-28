"""Gera o manual de uso / guia de apresentação em .docx.

Uso: python docs/build_manual.py
Gera: Manual_Simulador_Banqueiro.docx na raiz do projeto.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "imagens")
OUT_PATH = os.path.join(os.path.dirname(BASE_DIR), "Manual_Simulador_Banqueiro.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)

fig_counter = [0]


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def new_doc():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 18), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_para)

    return doc


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Simulador do Algoritmo do Banqueiro")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual de Uso e Guia de Apresentação")
    run.font.size = Pt(16)
    run.font.color.rgb = GRAY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Prevenção de deadlock por alocação de recursos: como o simulador funciona, "
        "como operá-lo e um roteiro pronto para apresentar em sala."
    )
    run.italic = True
    run.font.color.rgb = GRAY

    for _ in range(6):
        doc.add_paragraph()

    for label in ["Aluno(a): ______________________________________",
                  "Disciplina: ______________________________________",
                  "Data: 28 de agosto de 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label).font.color.rgb = GRAY

    doc.add_page_break()


def add_summary(doc, items):
    doc.add_heading("Sumário", level=1)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_image(doc, filename, caption, width_cm=15):
    fig_counter[0] += 1
    doc.add_picture(os.path.join(IMG_DIR, filename), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figura {fig_counter[0]} — {caption}")
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY


def add_table(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1F4E5F")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    run = p.add_run("P: " + question)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.add_run("R: " + answer)
    doc.add_paragraph()


# --------------------------------------------------------------------------

doc = new_doc()
add_title_page(doc)

add_summary(doc, [
    "1. Introdução",
    "2. Conceitos fundamentais (deadlock e o Algoritmo do Banqueiro)",
    "3. Visão geral do programa",
    "4. Guia de uso passo a passo",
    "5. Demonstração guiada dos três cenários prontos",
    "6. Roteiro sugerido para a apresentação",
    "7. Perguntas frequentes esperadas",
    "8. Detalhes técnicos do projeto",
    "9. Conclusão",
])

# 1. Introdução -------------------------------------------------------------
add_heading(doc, "1. Introdução")
add_para(
    doc,
    "Em um sistema operacional, vários processos disputam recursos limitados "
    "(memória, arquivos, impressoras, e assim por diante). Quando um processo "
    "retém um recurso e espera por outro que está retido por um segundo "
    "processo -- que por sua vez espera por algo retido pelo primeiro -- "
    "nenhum dos dois consegue prosseguir. Esse travamento permanente é "
    "chamado de deadlock (impasse)."
)
add_para(
    doc,
    "Este projeto implementa e visualiza o Algoritmo do Banqueiro, uma "
    "estratégia clássica (Dijkstra) para prevenir deadlocks: antes de "
    "conceder qualquer requisição de recurso, o algoritmo simula a alocação "
    "e verifica se ainda existe uma sequência em que todos os processos "
    "conseguem terminar. Se não existir, a requisição é negada."
)
add_para(
    doc,
    "Este documento serve para duas coisas: (1) ensinar a operar o "
    "simulador passo a passo, com telas reais do programa, e (2) fornecer "
    "um roteiro pronto -- com explicações, ordem sugerida e perguntas "
    "esperadas -- para apresentar o trabalho em sala de aula."
)

# 2. Conceitos fundamentais --------------------------------------------------
add_heading(doc, "2. Conceitos Fundamentais")

add_heading(doc, "2.1 A analogia do corredor estreito", level=2)
add_para(
    doc,
    "Imagine um corredor estreito da escola, largo o bastante para uma "
    "pessoa por vez. Dois grupos de amigos entram em direções opostas. Se "
    "cada grupo segurar sua posição esperando o outro \"dar espaço\", os "
    "dois ficam parados -- cada um esperando um recurso (o espaço livre do "
    "corredor) que só o outro pode liberar. Ninguém andou errado sozinho: o "
    "problema é a combinação das decisões de todos. É exatamente essa a "
    "situação que o Algoritmo do Banqueiro evita entre processos e recursos."
)

add_heading(doc, "2.2 As quatro condições necessárias para deadlock", level=2)
add_para(doc, "Um deadlock só ocorre se estas quatro condições estiverem presentes ao mesmo tempo:")
add_bullets(doc, [
    "Exclusão mútua -- o recurso só pode ser usado por um processo por vez.",
    "Posse e espera -- um processo mantém o que já tem enquanto espera por mais.",
    "Não preempção -- ninguém pode tomar à força um recurso de outro processo.",
    "Espera circular -- existe uma cadeia circular de processos, cada um esperando um recurso retido pelo próximo.",
], numbered=True)
add_para(
    doc,
    "O Algoritmo do Banqueiro ataca diretamente a condição de espera "
    "circular: ele nunca concede uma requisição que possa, no futuro, levar "
    "a essa cadeia fechada."
)

add_heading(doc, "2.3 Estado seguro, estado inseguro e deadlock", level=2)
add_bullets(doc, [
    "Estado seguro: existe pelo menos uma ordem em que todos os processos conseguem terminar.",
    "Estado inseguro: não há garantia dessa ordem -- ainda não é um deadlock, mas já é arriscado o suficiente para o Banqueiro recusar a próxima requisição perigosa.",
    "Deadlock: a espera circular já aconteceu de fato -- os processos envolvidos nunca mais progridem sozinhos.",
])

add_heading(doc, "2.4 As matrizes do algoritmo", level=2)
add_bullets(doc, [
    "Allocation -- quanto cada processo já está segurando de cada recurso.",
    "Max -- o máximo que cada processo pode vir a pedir (declarado no início).",
    "Need = Max − Allocation -- quanto cada processo ainda pode precisar.",
    "Available -- quanto sobrou livre de cada recurso.",
])
add_para(
    doc,
    "A cada requisição, o simulador testa, para cada processo ainda não "
    "finalizado, se Need <= Work (o que está livre no momento do teste). Se "
    "sim, aquele processo poderia terminar e devolver tudo que tem, "
    "aumentando o Work disponível para os próximos. Se, no final da "
    "varredura, todos os processos puderem ser encaixados nessa ordem, o "
    "estado é seguro."
)

add_image(doc, "13_ajuda.png", "Aba Ajuda do simulador, com a analogia do corredor e as 4 condições de deadlock.")

# 3. Visão geral --------------------------------------------------------------
doc.add_page_break()
add_heading(doc, "3. Visão Geral do Programa")
add_para(
    doc,
    "O simulador é um programa Python de janela única (Tkinter), com o "
    "grafo desenhado via matplotlib/networkx, organizado em 5 abas:"
)
add_table(doc, ["Aba", "O que faz"], [
    ["Configuração", "Cadastro manual de recursos e processos; inicia uma simulação do zero (Allocation = 0)."],
    ["Simulação", "Grafo de alocação de recursos, matrizes Allocation/Need/Available e formulário de requisição."],
    ["Log / Histórico", "Raciocínio passo a passo do algoritmo e uma linha do tempo com snapshots do estado."],
    ["Cenários prontos", "Três cenários pré-configurados para carregar com um clique (A, B e C)."],
    ["Ajuda", "A analogia do corredor e a explicação das condições de deadlock."],
], col_widths_cm=[4, 12])

add_heading(doc, "Como executar", level=2)
add_para(doc, "Dentro da pasta do projeto, com Python 3.10+ instalado:")
p = doc.add_paragraph()
run = p.add_run("python main.py")
run.font.name = "Consolas"
run.font.size = Pt(11)
add_para(
    doc,
    "É necessário ter as bibliotecas matplotlib e networkx instaladas "
    "(tkinter já vem com o Python no Windows). Para rodar os testes "
    "automatizados do algoritmo:"
)
p = doc.add_paragraph()
run = p.add_run("python -m unittest discover -s tests -v")
run.font.name = "Consolas"
run.font.size = Pt(11)

# 4. Guia de uso ----------------------------------------------------------
doc.add_page_break()
add_heading(doc, "4. Guia de Uso Passo a Passo")

add_heading(doc, "4.1 Aba Configuração -- montando uma simulação do zero", level=2)
add_para(doc, "A aba começa vazia, com um formulário para recursos e outro para processos:")
add_image(doc, "01_config_vazia.png", "Aba Configuração, ainda vazia.")

add_para(doc, "1. Cadastre um recurso: digite um nome e a quantidade total de instâncias, e clique em \"Adicionar Recurso\".")
add_image(doc, "02_config_recursos.png", "Dois recursos (A e B) já cadastrados.")

add_para(doc, "2. Cadastre um processo: digite um nome e, para cada recurso já cadastrado, o máximo que esse processo pode vir a pedir. Clique em \"Adicionar Processo\".")
add_image(doc, "03_config_processos.png", "Dois processos (P0 e P1) já cadastrados, cada um com seu vetor Max.")

add_para(doc, "3. Clique em \"Iniciar Simulação (Allocation = 0)\". O programa monta o estado inicial -- todos os processos começam sem nenhum recurso alocado -- e leva você para a aba Simulação.")
add_image(doc, "04_simulacao_manual_iniciada.png", "Simulação recém-iniciada: Allocation zerada, grafo sem nenhuma seta ainda.")

add_heading(doc, "4.2 Aba Simulação -- o grafo, as matrizes e o formulário de requisição", level=2)
add_para(doc, "É a tela principal, dividida em duas partes:")
add_bullets(doc, [
    "À esquerda: o Grafo de Alocação de Recursos. Círculos são processos, quadrados são recursos (com \"livre/total\" escrito dentro). Setas pretas cheias indicam alocação (recurso -> processo); setas vermelhas tracejadas indicam um pedido pendente ou negado (processo -> recurso).",
    "À direita: as tabelas Allocation, Need e Available, sempre atualizadas, e o formulário para escolher um processo, digitar quanto de cada recurso ele está pedindo, e clicar em \"Solicitar\".",
])
add_para(
    doc,
    "As cores dos processos também têm significado: verde é pronto/executando, "
    "amarelo é bloqueado (esperando um recurso) e cinza é terminado."
)
add_image(doc, "05_simulacao_manual_pedido.png", "Depois de um pedido concedido: a Allocation do processo escolhido aumentou e o Available diminuiu.")
add_para(
    doc,
    "Repare também nos dois botões de rádio \"Protegido (com Banqueiro)\" e "
    "\"Livre (sem prevenção)\": eles alternam se o simulador vai verificar "
    "segurança antes de conceder cada pedido ou não -- isso é essencial para "
    "o Cenário C, explicado adiante."
)

add_heading(doc, "4.3 Aba Log / Histórico -- o raciocínio do algoritmo", level=2)
add_para(
    doc,
    "Toda vez que algo acontece (um pedido, uma negação, uma detecção de "
    "deadlock), o programa registra cada etapa do raciocínio nessa aba: "
    "quais comparações Need <= Work foram feitas, se cada processo pôde "
    "prosseguir, e a conclusão final. Do lado direito fica a linha do "
    "tempo (timeline), com um resumo de cada evento -- clicando em um "
    "deles, é possível ver como as matrizes estavam naquele momento."
)

add_heading(doc, "4.4 Aba Cenários prontos", level=2)
add_para(
    doc,
    "A forma mais rápida de demonstrar o simulador: três botões, cada um "
    "carregando uma situação já configurada (explicados em detalhe na "
    "seção 5)."
)

add_heading(doc, "4.5 Aba Ajuda", level=2)
add_para(doc, "Contém a analogia do corredor e a explicação das quatro condições de deadlock, útil para consultar durante a apresentação (ver Figura 1).")

# 5. Cenários ----------------------------------------------------------------
doc.add_page_break()
add_heading(doc, "5. Demonstração Guiada dos Três Cenários")
add_para(
    doc,
    "Esta é a parte mais importante para a apresentação: os três cenários "
    "prontos mostram, na prática, um estado seguro, uma negação e um "
    "deadlock real."
)

add_heading(doc, "5.1 Cenário A -- Estado seguro (exemplo clássico)", level=2)
add_para(
    doc,
    "5 processos disputando 3 tipos de recursos (A=10, B=5, C=7). Ao "
    "carregar, o estado já é seguro: existe uma sequência de execução -- "
    "<P1, P3, P4, P0, P2> -- em que todos os processos conseguem terminar "
    "sem travar. Esse é o exemplo clássico do livro de Sistemas "
    "Operacionais (Silberschatz)."
)
add_image(doc, "06_cenario_a.png", "Cenário A carregado: grafo, matrizes Allocation/Need/Available e status \"pronto\" para todos os processos.")
add_para(doc, "O que dizer na apresentação: mesmo com vários processos precisando de bastante recurso (veja a coluna Need), sempre existe alguém que pode terminar primeiro e devolver o que usou, destravando os próximos -- por isso o estado é seguro.")

add_heading(doc, "5.2 Cenário B -- Requisição negada (estado inseguro)", level=2)
add_para(
    doc,
    "3 processos disputando 2 recursos (A=3, B=3). O estado inicial é "
    "seguro, mas a descrição do cenário sugere um pedido perigoso: o "
    "processo P1 solicitando (1, 0)."
)
add_image(doc, "07_cenario_b_inicial.png", "Cenário B logo após carregar -- ainda em estado seguro.")
add_para(doc, "Ao ir na aba Simulação, escolher P1, digitar 1 no primeiro recurso e 0 no segundo, e clicar em Solicitar, o pedido é negado:")
add_image(doc, "08_cenario_b_negado.png", "Pedido de P1 negado: a Allocation voltou exatamente ao que era antes (rollback), P1 aparece bloqueado (amarelo) com uma seta vermelha tracejada pendente.")
add_para(doc, "Na aba Log, é possível mostrar o motivo exato:")
add_image(doc, "09_cenario_b_log.png", "Log do Cenário B: cada comparação Need <= Work aparece, terminando em \"Estado INSEGURO\" e \"ROLLBACK\".")
add_para(doc, "O que dizer na apresentação: mesmo havendo recurso disponível na hora do pedido, o Banqueiro enxerga à frente e percebe que, se conceder, ninguém mais conseguiria progredir depois -- por isso prefere negar e desfazer a alocação (rollback) a arriscar um travamento.")

add_heading(doc, "5.3 Cenário C -- Deadlock real por espera circular (modo Livre)", level=2)
add_para(
    doc,
    "3 processos e 3 recursos de instância única, cada um já alocado por "
    "um processo diferente. Esse cenário carrega automaticamente em modo "
    "Livre -- ou seja, sem a proteção do Banqueiro -- exatamente para "
    "mostrar o que aconteceria sem o algoritmo."
)
add_image(doc, "10_cenario_c_inicial.png", "Cenário C recém-carregado: cada processo já segura um recurso, Available = 0 em tudo.")
add_para(
    doc,
    "Seguindo os pedidos sugeridos na descrição do cenário (P0 pede o "
    "recurso de P1, P1 pede o de P2, e P2 pede o de P0), as três "
    "requisições ficam pendentes e fecham um ciclo:"
)
add_image(doc, "11_cenario_c_ciclo.png", "As três requisições fecham um ciclo -- destacado em vermelho no grafo -- e o título já avisa \"CICLO DETECTADO\".")
add_para(doc, "Clicando em \"Verificar Deadlock\", o algoritmo de detecção confirma formalmente o travamento:")
add_image(doc, "12_cenario_c_log.png", "Log de detecção: cada processo é testado contra o que está pendente, e nenhum consegue progredir -- \"DEADLOCK detectado\".")
add_para(
    doc,
    "O que dizer na apresentação: esse é o contraste-chave do trabalho -- "
    "se esse mesmo cenário rodasse em modo Protegido, a segunda ou terceira "
    "requisição já teria sido negada antes de o ciclo se fechar, exatamente "
    "como no Cenário B."
)

# 6. Roteiro -------------------------------------------------------------------
doc.add_page_break()
add_heading(doc, "6. Roteiro Sugerido para a Apresentação")
add_para(doc, "Um roteiro de aproximadamente 8 a 10 minutos, usando só a aba Cenários prontos e a de Log:")
add_table(doc, ["Tempo", "O que fazer", "O que falar"], [
    ["~1 min", "Abrir na aba Ajuda", "Contar a analogia do corredor e citar as 4 condições de deadlock."],
    ["~2 min", "Carregar Cenário A", "Explicar Allocation/Max/Need/Available e por que esse estado é seguro."],
    ["~2 min", "Carregar Cenário B e fazer o pedido de P1", "Mostrar a negação no grafo e depois o raciocínio linha a linha no Log."],
    ["~3 min", "Carregar Cenário C, fazer os 3 pedidos e clicar em Verificar Deadlock", "Mostrar o ciclo em vermelho e a confirmação no Log; comparar com o que aconteceria em modo Protegido."],
    ["~1 min", "Encerramento", "Resumir: o Banqueiro nunca deixa o sistema chegar a um deadlock, recusando qualquer pedido que levaria a um estado inseguro."],
], col_widths_cm=[2.2, 6.3, 7.5])
add_para(doc, "Dica: pratique a sequência de cliques antes da apresentação (Cenários prontos -> Simulação -> Log) para não perder tempo procurando os botões na hora.")

# 7. Perguntas frequentes -------------------------------------------------------
doc.add_page_break()
add_heading(doc, "7. Perguntas Frequentes Esperadas")
add_qa(doc,
    "Por que o Banqueiro nega um pedido se, naquele instante, existe recurso disponível?",
    "Porque \"ter disponível agora\" não é suficiente -- o algoritmo simula a concessão e verifica se, depois dela, ainda existe alguma ordem em que todos os processos conseguem terminar. Se não existir, é arriscado demais conceder.")
add_qa(doc,
    "Qual a diferença entre estado inseguro e deadlock?",
    "Estado inseguro é quando não há garantia de uma sequência segura -- pode ou não terminar em deadlock, dependendo do que cada processo pedir depois. Deadlock é quando a espera circular já aconteceu de fato e ninguém mais progride sozinho.")
add_qa(doc,
    "Como o simulador consegue mostrar um deadlock de verdade, se o Banqueiro sempre previne?",
    "Existe um modo \"Livre\", usado no Cenário C, que desliga a verificação de segurança de propósito, só para fins didáticos -- assim é possível ver a espera circular se formando e sendo confirmada por um algoritmo de detecção separado.")
add_qa(doc,
    "O que acontece quando uma requisição é negada -- os recursos ficam perdidos?",
    "Não. O simulador testa a alocação, e se o resultado for inseguro, desfaz exatamente essa tentativa (rollback) -- as matrizes voltam a ser idênticas ao estado anterior ao pedido.")
add_qa(doc,
    "Esse algoritmo é usado em sistemas operacionais reais?",
    "O Algoritmo do Banqueiro é mais um modelo teórico/didático (proposto por Edsger Dijkstra) do que uma solução usada diretamente em SOs de propósito geral, porque exige saber de antemão o máximo que cada processo vai precisar -- algo raramente conhecido na prática. Sistemas reais tendem a usar outras estratégias, como detecção e recuperação, ou simplesmente evitar ciclos de espera por convenção (ex.: sempre pedir recursos em uma mesma ordem).")
add_qa(doc,
    "Por que as matrizes usam o vetor Need em vez de Max diretamente?",
    "Need = Max − Allocation representa o que falta, que é exatamente o que interessa para saber se um processo consegue terminar com o que ainda está disponível (Work).")

# 8. Detalhes técnicos ------------------------------------------------------
doc.add_page_break()
add_heading(doc, "8. Detalhes Técnicos do Projeto")
add_para(doc, "Para quem quiser (ou for perguntado sobre) a implementação:")
add_bullets(doc, [
    "Linguagem: Python 3, interface gráfica com Tkinter.",
    "Grafo de alocação: matplotlib + networkx, embutido na janela.",
    "core/ -- lógica pura do algoritmo (Safety Algorithm, Resource-Request Algorithm, detecção de deadlock), sem nenhuma dependência de interface gráfica.",
    "simulation/ -- motor de simulação: modos Protegido e Livre, histórico de eventos e snapshots.",
    "scenarios/ -- definição dos três cenários prontos (A, B e C).",
    "gui/ -- as cinco abas da interface.",
    "tests/ -- 33 testes automatizados cobrindo o algoritmo (incluindo o caso clássico do livro-texto e o teste de rollback exato).",
])
p = doc.add_paragraph()
run = p.add_run("python -m unittest discover -s tests -v")
run.font.name = "Consolas"
run.font.size = Pt(11)
add_para(doc, "roda toda a suíte de testes e mostra cada verificação passando.")

# 9. Conclusão ------------------------------------------------------------------
doc.add_page_break()
add_heading(doc, "9. Conclusão")
add_para(
    doc,
    "O simulador cobre, de ponta a ponta, o ciclo pedido pelo trabalho: "
    "simular múltiplos processos disputando recursos, implementar o "
    "Algoritmo do Banqueiro completo (incluindo o Safety Algorithm), "
    "visualizar o grafo de alocação, representar as matrizes de estado, "
    "demonstrar cenários que levam a deadlock e como são prevenidos, e "
    "registrar em log o raciocínio de cada decisão."
)
add_para(
    doc,
    "Para a apresentação, o caminho mais direto é: Ajuda -> Cenário A -> "
    "Cenário B (com o pedido negado) -> Cenário C (com o deadlock real) -> "
    "Log. Esse percurso, sozinho, já conta toda a história do algoritmo."
)

doc.save(OUT_PATH)
print("Documento salvo em:", OUT_PATH)
